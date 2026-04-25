"""
ingestion/ingest.py
────────────────────────────────────────────────────────────────────────────────
Pipeline: read corpus → extract text (PyMuPDF for PDFs, python-docx for DOCX,
plain read for Markdown) → chunk → embed → upsert to Qdrant.

Re-running is idempotent: documents are keyed by (doc_id, chunk_index).
"""
from __future__ import annotations

import hashlib
import json
import logging
import re
import sys
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF
import docx as python_docx
from rank_bm25 import BM25Okapi
from qdrant_client import QdrantClient
from qdrant_client.models import (
    Distance,
    VectorParams,
    PointStruct,
    Filter,
    FieldCondition,
    MatchValue,
)

sys.path.insert(0, str(Path(__file__).parent.parent))
import config

logger = logging.getLogger(__name__)

# ─── Constants ────────────────────────────────────────────────────────────────
CHUNK_SIZE = 600       # tokens / words (approximate)
CHUNK_OVERLAP = 80

# ─── Sentence-transformers embeddings ────────────────────────────────────────
_embed_model = None


def _get_embed_model():
    global _embed_model
    if _embed_model is None:
        from sentence_transformers import SentenceTransformer
        logger.info("Loading embedding model: %s", config.EMBEDDING_MODEL)
        _embed_model = SentenceTransformer(config.EMBEDDING_MODEL)
    return _embed_model


def embed(texts: list[str]) -> list[list[float]]:
    return _get_embed_model().encode(texts, normalize_embeddings=True).tolist()


def embed_dim() -> int:
    return _get_embed_model().get_sentence_embedding_dimension()


# ─── Qdrant client singleton ──────────────────────────────────────────────────
_qdrant: QdrantClient | None = None


def get_qdrant() -> QdrantClient:
    global _qdrant
    if _qdrant is None:
        if config.QDRANT_URL:
            _qdrant = QdrantClient(url=config.QDRANT_URL)
        else:
            _qdrant = QdrantClient(":memory:")
    return _qdrant


def ensure_collection(dim: int = 0) -> None:
    if dim == 0:
        dim = embed_dim()
    client = get_qdrant()
    existing = [c.name for c in client.get_collections().collections]
    if config.QDRANT_COLLECTION not in existing:
        client.create_collection(
            collection_name=config.QDRANT_COLLECTION,
            vectors_config=VectorParams(size=dim, distance=Distance.COSINE),
        )
        logger.info("Created Qdrant collection '%s'", config.QDRANT_COLLECTION)


# ─── Text extraction ──────────────────────────────────────────────────────────
def extract_pdf_to_markdown(path: Path) -> str:
    """Extract PDF to markdown using PyMuPDF."""
    doc = fitz.open(str(path))
    pages: list[str] = []
    for page_num, page in enumerate(doc, start=1):
        blocks = page.get_text("dict")["blocks"]
        lines: list[str] = []
        for block in blocks:
            if block["type"] != 0:  # skip images
                continue
            for line in block["lines"]:
                text = " ".join(span["text"] for span in line["spans"]).strip()
                if text:
                    # heuristic heading detection (large/bold font)
                    spans = line["spans"]
                    if spans and (spans[0].get("size", 12) > 13 or "Bold" in spans[0].get("font", "")):
                        lines.append(f"## {text}")
                    else:
                        lines.append(text)
        if lines:
            pages.append(f"<!-- page {page_num} -->\n" + "\n".join(lines))
    doc.close()
    return "\n\n".join(pages)


def extract_docx_to_markdown(path: Path) -> str:
    """Extract DOCX to markdown using python-docx."""
    doc = python_docx.Document(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name.lower()
        if "heading 1" in style:
            lines.append(f"# {text}")
        elif "heading 2" in style:
            lines.append(f"## {text}")
        elif "heading 3" in style:
            lines.append(f"### {text}")
        else:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    return "\n\n".join(lines)


def extract_markdown(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


# ─── Chunking ─────────────────────────────────────────────────────────────────
def _word_chunks(text: str, size: int, overlap: int) -> list[str]:
    words = text.split()
    chunks, i = [], 0
    while i < len(words):
        chunk = words[i : i + size]
        chunks.append(" ".join(chunk))
        i += size - overlap
    return [c for c in chunks if c.strip()]


def chunk_document(doc_id: str, text: str, metadata: dict) -> list[dict]:
    raw_chunks = _word_chunks(text, CHUNK_SIZE, CHUNK_OVERLAP)
    result = []
    for idx, chunk in enumerate(raw_chunks):
        result.append(
            {
                "chunk_id": f"{doc_id}__chunk{idx}",
                "doc_id": doc_id,
                "chunk_index": idx,
                "text": chunk,
                "metadata": metadata,
            }
        )
    return result


# ─── BM25 index (in-process, rebuilt from Qdrant payload on load) ─────────────
class BM25Index:
    """Lightweight BM25 wrapper that operates on the same chunk corpus."""

    def __init__(self) -> None:
        self._chunks: list[dict] = []
        self._bm25: BM25Okapi | None = None

    def build(self, chunks: list[dict]) -> None:
        self._chunks = chunks
        tokenised = [c["text"].lower().split() for c in chunks]
        self._bm25 = BM25Okapi(tokenised)

    def query(self, query: str, top_k: int) -> list[tuple[dict, float]]:
        if self._bm25 is None or not self._chunks:
            return []
        scores = self._bm25.get_scores(query.lower().split())
        ranked = sorted(enumerate(scores), key=lambda x: x[1], reverse=True)[:top_k]
        return [(self._chunks[i], float(s)) for i, s in ranked if s > 0]

    def size(self) -> int:
        return len(self._chunks)


_bm25_index = BM25Index()


def get_bm25_index() -> BM25Index:
    return _bm25_index


# ─── Main ingestion function ───────────────────────────────────────────────────
def ingest_corpus(corpus_dir: Path, metadata_file: Path) -> int:
    """
    Ingest all documents in corpus_dir into Qdrant + BM25.
    Returns total number of chunks indexed.
    """
    with open(metadata_file) as f:
        meta_json: dict = json.load(f)

    docs_meta: dict[str, dict] = meta_json.get("documents", {})
    all_chunks: list[dict] = []
    skipped = 0

    for doc_id, doc_meta in docs_meta.items():
        fmt = doc_meta.get("format", "markdown")

        # locate file
        candidates = list(corpus_dir.glob(f"{doc_id}.*")) + list(
            corpus_dir.glob(f"**/{doc_id}.*")
        )
        if not candidates:
            # try matching by title slug
            candidates = list(corpus_dir.glob(f"*{doc_id}*"))
        if not candidates:
            logger.warning("No file found for %s — skipping", doc_id)
            skipped += 1
            continue

        path = candidates[0]
        logger.info("Ingesting %s (%s) from %s", doc_id, fmt, path.name)

        try:
            if fmt == "pdf" or path.suffix.lower() == ".pdf":
                text = extract_pdf_to_markdown(path)
            elif fmt == "docx" or path.suffix.lower() == ".docx":
                text = extract_docx_to_markdown(path)
            else:
                text = extract_markdown(path)
        except Exception as exc:
            logger.error("Failed to extract %s: %s", doc_id, exc)
            skipped += 1
            continue

        if not text.strip():
            logger.warning("Empty extraction for %s", doc_id)
            skipped += 1
            continue

        chunks = chunk_document(doc_id, text, doc_meta)
        all_chunks.extend(chunks)

    if not all_chunks:
        logger.error("No chunks produced — check corpus directory.")
        return 0

    # ── Embed ────────────────────────────────────────────────────────────────
    texts = [c["text"] for c in all_chunks]
    logger.info("Embedding %d chunks …", len(texts))
    batch_size = 64
    vectors: list[list[float]] = []
    for i in range(0, len(texts), batch_size):
        vectors.extend(embed(texts[i : i + batch_size]))

    dim = len(vectors[0])
    ensure_collection(dim)

    # ── Upsert to Qdrant ─────────────────────────────────────────────────────
    client = get_qdrant()
    points: list[PointStruct] = []
    for chunk, vec in zip(all_chunks, vectors):
        uid = int(
            hashlib.md5(chunk["chunk_id"].encode()).hexdigest()[:16], 16
        ) % (2**63)
        points.append(
            PointStruct(
                id=uid,
                vector=vec,
                payload={
                    "chunk_id": chunk["chunk_id"],
                    "doc_id": chunk["doc_id"],
                    "chunk_index": chunk["chunk_index"],
                    "text": chunk["text"],
                    **chunk["metadata"],
                },
            )
        )

    batch_size_qdrant = 128
    for i in range(0, len(points), batch_size_qdrant):
        client.upsert(
            collection_name=config.QDRANT_COLLECTION,
            points=points[i : i + batch_size_qdrant],
            wait=True,
        )

    # ── Build BM25 ───────────────────────────────────────────────────────────
    _bm25_index.build(all_chunks)

    logger.info(
        "Ingestion complete: %d chunks from %d docs (%d skipped).",
        len(all_chunks),
        len(docs_meta) - skipped,
        skipped,
    )
    return len(all_chunks)


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    n = ingest_corpus(config.CORPUS_DIR, config.METADATA_FILE)
    print(f"Indexed {n} chunks.")